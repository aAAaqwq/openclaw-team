# -*- coding: utf-8 -*-
"""
DOCX文档美化脚本模板
用于将Word文档中的内容转换为专业表格格式
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import glob
import re

# ============ 配置区域 ============
INPUT_FILE = r"输入文档路径.docx"
OUTPUT_DIR = r"输出目录"
SECTION_TITLE = "章节标题"  # 如 "3.1  实体识别"
# ================================

# ============ 样式函数 ============
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), edge_data.get('val', 'single'))
            border.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            border.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_background(cell, color="4472C4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def format_table(table):
    """应用表格样式"""
    border_props = {'val': 'single', 'sz': 4, 'color': '000000'}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_props, left=border_props,
                          bottom=border_props, right=border_props)
    for cell in table.rows[0].cells:
        set_cell_background(cell, "4472C4")
        for paragraph in cell.paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(table.rows[1:], 1):
        bg_color = "E7E6E6" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def create_table(doc, headers, rows):
    """创建表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row_data in enumerate(rows):
        for j, text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = text
    format_table(table)
    return table

def get_next_version(path):
    """获取下一个版本号"""
    files = glob.glob(path + '/*v0.*.docx') + glob.glob(path + '/*v1.*.docx')
    versions = []
    for f in files:
        m = re.search(r'v(\d+)\.(\d+)', f)
        if m:
            versions.append((int(m.group(1)), int(m.group(2))))
    if not versions:
        return 'v0.2'
    versions.sort()
    last = versions[-1]
    if last[1] < 9:
        return f'v{last[0]}.{last[1] + 1}'
    else:
        return f'v{last[0] + 1}.0'

# ============ 数据定义（根据需要修改）============
TABLE_DATA = [
    ["列1", "列2", "列3"],
    ["数据1", "数据2", "数据3"],
    ["数据4", "数据5", "数据6"],
]

# ============ 主流程 ============
def main():
    # 读取文档
    doc = Document(INPUT_FILE)

    # 获取版本号
    version = get_next_version(OUTPUT_DIR)
    print(f"版本号: {version}")

    # 找到章节并替换
    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        if SECTION_TITLE in para.text:
            print(f"找到章节: {SECTION_TITLE}")

            # 创建表格
            table = create_table(doc, TABLE_DATA[0], TABLE_DATA[1:])

            # 插入表格
            table._element.getparent().insert(
                para._element.getparent().index(para._element) + 1,
                table._element
            )
            break

    # 保存
    output_path = f"{OUTPUT_DIR}/文档-{version}.docx"
    doc.save(output_path)
    print(f"保存到: {output_path}")

if __name__ == "__main__":
    main()
